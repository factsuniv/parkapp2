#!/usr/bin/env python3
"""
Extract content from Word document (.docx)
"""

from docx import Document
import sys
import os

def extract_docx_content(file_path):
    """Extract text content from a Word document"""
    try:
        # Load the document
        doc = Document(file_path)
        
        # Extract all paragraphs
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                content.append(paragraph.text.strip())
        
        # Extract content from tables if any
        for table in doc.tables:
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    content.append(" | ".join(row_content))
        
        return "\n\n".join(content)
    
    except Exception as e:
        return f"Error extracting content: {str(e)}"

if __name__ == "__main__":
    file_path = "/workspace/user_input_files/Park .docx"
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)
    
    content = extract_docx_content(file_path)
    print("=== EXTRACTED CONTENT ===")
    print(content)
    print("=== END CONTENT ===")
    
    # Save to a text file for easier access
    with open("/workspace/extracted_content/park_business_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"\nContent saved to: /workspace/extracted_content/park_business_content.txt")